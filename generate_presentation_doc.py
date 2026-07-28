"""Generates the AI Financial Analyst System presentation document as DOCX."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"c:\Users\nithi\OneDrive\Desktop\AI Financial Analyst System - Full Presentation.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)   # blue
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x11, 0x39, 0x7D)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after   = Pt(3)
    return p

def blockquote(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic    = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    # left border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "4A90D9")
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def add_table(headers, rows, header_color="1A56DB", alt_color="EBF5FB"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = alt_color if ri % 2 == 1 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    return t

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 90)
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("AI Financial Analyst System")
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Full Picture — What We Built, What We Are Building, and Where We Are Going")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Prepared for Internal Presentation  |  July 2026")
r3.italic = True; r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — What Is This System?
# ══════════════════════════════════════════════════════════════════════════════
heading1("SECTION 1 — What Is This System?")
divider()

body("This is an AI-powered trading assistant that does three things:")
bullet("Studies the past — reads months or years of stock price data")
bullet("Predicts the future — decides: BUY, SELL, or HOLD")
bullet("Checks itself — compares every prediction against what actually happened")

doc.add_paragraph()
body("Simple Analogy:", bold=True)
blockquote(
    "Think of it like a cricket analyst. Before a match, the analyst studies 3 years of data — "
    "player form, pitch conditions, weather — and predicts the score. After the match, we check: "
    "was the analyst right? Over time, the analyst gets smarter because we feed it its own mistakes."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — What Is Already Built
# ══════════════════════════════════════════════════════════════════════════════
heading1("SECTION 2 — What Is Already Built and Working")
divider()

add_table(
    ["Feature", "Status", "What It Does"],
    [
        ["Stock Prediction",       "WORKING", "Predicts BUY/SELL/HOLD for any stock"],
        ["Options Backtest",       "WORKING", "Tests options strategies (calls/puts) against history"],
        ["Accuracy Tracking",      "WORKING", "Logs every prediction, marks right/wrong"],
        ["Calibration System",     "WORKING", "Tells AI its own past mistakes per stock"],
        ["Market Movers Feed",     "WORKING", "Feeds today's top gaining/losing stocks to AI"],
        ["Earnings Calendar",      "WORKING", "Warns AI about upcoming earnings before prediction"],
        ["Exit Rules",             "FIXED",   "Take profit / stop loss now properly wired"],
        ["567 Automated Tests",    "PASSING", "Every feature verified, zero crashes"],
    ],
    header_color="1A56DB",
    alt_color="EBF5FB",
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — The Main Problem
# ══════════════════════════════════════════════════════════════════════════════
heading1("SECTION 3 — The Main Problem Right Now")
divider()

body("The AI is right only 24 times out of 100.", bold=True)
doc.add_paragraph()
body("That means 76% of the time, the AI makes the wrong call on direction (BUY/SELL/HOLD).")

doc.add_paragraph()
body("Why is this happening?", bold=True)
body("The AI says HOLD too often — even when the market is clearly going up or down.")

doc.add_paragraph()
body("Real Example from our data:", bold=True)
blockquote(
    "Tesla dropped sharply from $420 to $371.\n"
    "The AI said: HOLD  (wrong)\n"
    "It should have said: SELL\n"
    "The AI played it safe instead of taking a clear position."
)

body("What we already fixed to reduce this:", bold=True)
bullet("Lowered the threshold that forces HOLD — AI now takes more decisive positions")
bullet("Fixed the AI reading its own calibration data (it was silently failing before)")
bullet("AI now receives today's market movers before making a decision")

doc.add_paragraph()
body("Target after all fixes: 50%+ accuracy on SPY and TSLA", bold=True, color=(0x06, 0x6A, 0x2D))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — What We Are Building Next
# ══════════════════════════════════════════════════════════════════════════════
heading1("SECTION 4 — What We Are Building Next (In Order)")
divider()

# --- Build 1 ---
heading2("BUILD 1 — Accuracy Improvement on SPY and TSLA")
body("Problem:", bold=True)
body("24% accuracy is not reliable enough to use in real trading.")

doc.add_paragraph()
body("What we build:", bold=True)
body("Multi-agent approach — Instead of one AI making the decision, we split into 3 mini-analysts:")
bullet("Analyst 1 — reads technical signals (RSI, MACD, Bollinger Bands)")
bullet("Analyst 2 — reads market mood and sector trends")
bullet("Analyst 3 — reads recent news and events")
bullet("All 3 vote → majority wins")

doc.add_paragraph()
body("We also feed every AI its past mistakes per stock:")
blockquote('"For TSLA specifically, you were wrong 76% of the time. Here is what you missed."')

body("Analogy:", bold=True)
blockquote(
    "Instead of one analyst making the call, it is a committee of 3. "
    "One reads the charts, one reads the news, one reads the market mood. "
    "They debate and vote. That gives a better answer than one person alone."
)
body("Goal: 50%+ decision accuracy on SPY and TSLA within next sprint.", bold=True, color=(0x06, 0x6A, 0x2D))

divider()

# --- Build 2 ---
heading2("BUILD 2 — Bulk Testing Across Top 10 S&P 500 Stocks")
body("Problem:", bold=True)
body("We test one stock at a time, manually. It takes hours.")

doc.add_paragraph()
body("What we build:", bold=True)
bullet("One-click runner that tests all 10 stocks: AAPL, MSFT, AMZN, GOOGL, TSLA, NVDA, META, BRK, JPM, UNH")
bullet("Tests across 6 different time windows over the past 3 years")
bullet("Shows a single summary table automatically")

add_table(
    ["Symbol", "Accuracy", "HOLD Rate", "Avg Return Error", "Best / Worst Period"],
    [
        ["SPY",  "54%", "22%", "1.2%", "Strong in 2024 Q1"],
        ["TSLA", "31%", "48%", "4.7%", "Weak all periods"],
        ["NVDA", "61%", "18%", "0.9%", "Consistently strong"],
    ],
)

body("Analogy:", bold=True)
blockquote(
    "Like a school report card — not just one exam, but every subject, every term. "
    "You see exactly where the student is strong and where they are weak."
)
body("Goal: Complete accuracy picture across all 10 stocks in one overnight run.", bold=True, color=(0x06, 0x6A, 0x2D))

divider()

# --- Build 3 ---
heading2("BUILD 3 — Free-Form Natural Language Input")
body("Problem:", bold=True)
body("The system today only accepts a strict form. Real traders don't type like that.")

doc.add_paragraph()
body("What the system accepts today:")
blockquote("Symbol: SPY | Start: 2026-01-01 | Capital: $50,000")

body("What real traders type:")
blockquote(
    '"Bought SPX 7570 call, premium 1.40"\n'
    '"I think Tesla is going to drop this week"\n'
    '"SPY 450 put, 7 days out"'
)

body("What we build:", bold=True)
bullet("A natural language layer that reads any sentence")
bullet("Extracts: symbol, direction, option type, strike, premium automatically")
bullet("Uses Gemini AI — same way ChatGPT reads any sentence")

doc.add_paragraph()
body('Example:', bold=True)
blockquote(
    'Input: "Buy SPX 7570 call for next Friday"\n'
    'System reads: Symbol=SPX | Action=BUY | Type=Call | Strike=7570 | Expiry=Friday'
)

body("Analogy:", bold=True)
blockquote(
    "Old system = new employee who only follows exact written instructions.\n"
    "New system = smart employee who understands what you mean even if you say it casually."
)

divider()
doc.add_page_break()

# --- Build 4 ---
heading2("BUILD 4 — Discord Group Integration (Live Deployment)")
body("Problem:", bold=True)
body(
    "Signals are already flowing in Discord groups all day. But the AI is not reading them. "
    "Someone has to manually copy-paste into the system."
)

doc.add_paragraph()
body("What we build:", bold=True)
bullet("A bot that sits in your Discord trading group 24/7")
bullet("Reads every message as it arrives")
bullet("When it detects a trading signal, runs the prediction automatically")
bullet("Posts the result back into the channel within seconds")

doc.add_paragraph()
body("Example of how it works in Discord:", bold=True)
blockquote(
    "Person A: Bought SPX 7570 call, premium 1.40\n\n"
    "AI Bot: Signal detected. Running analysis...\n"
    "SPX 7570 Call | AI Decision: HOLD | Confidence: 61%\n"
    "Historical win rate for this setup (7 DTE): 43% in similar conditions. Proceed with caution."
)

body("Analogy:", bold=True)
blockquote(
    "Like having a silent analyst sitting in your group chat all day. "
    "The moment someone says 'buying Tesla calls,' the analyst instantly pulls up the charts, "
    "runs the model, and posts the result back in 10 seconds."
)

divider()

# --- Build 5 ---
heading2("BUILD 5 — Full Day Conversation Memory")
body("Problem:", bold=True)
body("The bot today has no memory. Every message is treated as if nothing was said before.")

doc.add_paragraph()
body("Real scenario that breaks today:", bold=True)
blockquote(
    "9:00 AM — Person A: 'Bought SPX 7570 calls'\n"
    "2:30 PM — Person A: 'Sold it off for a small loss'\n\n"
    "The bot at 2:30 PM has no idea what 'it' refers to. It cannot process this."
)

body("What we build:", bold=True)
bullet("The bot holds the entire day's conversation in memory")
bullet("Every message stored with: who said it, what time, what trade it refers to")
bullet("Memory resets at midnight — clean slate every morning, full memory all day")

body("Analogy:", bold=True)
blockquote(
    "Old bot = colleague with amnesia. Every call, you explain everything from scratch.\n"
    "New bot = colleague with a notebook. At 3pm, they remember exactly what you discussed at 9am."
)

divider()

# --- Build 6 ---
heading2("BUILD 6 — Reply Chain Understanding")
body("Problem:", bold=True)
body("When someone replies to a specific message, the bot does not know which message they are replying to.")

blockquote(
    "Person A: Bought SPX 7570 call\n"
    "Person B (replying to A): I sold that off already\n\n"
    "Today the bot reads Person B's message and has no idea what 'that' means. It is lost."
)

body("What we build:", bold=True)
bullet("Store message IDs and reply-to IDs alongside every message")
bullet("When processing any reply, the bot attaches the original message automatically")
bullet("The AI sees both messages together — the original trade AND the reply")

doc.add_paragraph()
body(
    "Result: The bot now understands: 'I sold that off' = "
    "Person B exited the SPX 7570 call that Person A entered.",
    bold=True, color=(0x06, 0x6A, 0x2D)
)

divider()
doc.add_page_break()

# --- Build 7 ---
heading2("BUILD 7 — Signal Performance Tracking (Who Is Actually Right?)")
body("Problem:", bold=True)
body(
    "In any trading group, some people give great signals, some give terrible ones. "
    "Right now there is no way to know who to trust."
)

doc.add_paragraph()
body("What we build:", bold=True)
bullet("Every signal detected in the group is logged: who said it, what trade, at what time")
bullet("When the trade window closes, system checks: did the price move in the predicted direction?")
bullet("Shows a live leaderboard:")

add_table(
    ["Person", "Signals Given", "Correct", "Accuracy", "Best Asset"],
    [
        ["Person A", "14", "10", "71%", "SPX"],
        ["Person B", "8",  "4",  "50%", "TSLA"],
        ["Person C", "5",  "1",  "20%", "NVDA"],
    ],
)

body("Also scores groups:", bold=True)
blockquote("Options Study Discord — 65% signal accuracy  vs  Random Group — 38%")

body("Analogy:", bold=True)
blockquote(
    "Like tracking a cricket player's batting average across a full season. "
    "After 50 innings you know exactly who to trust at the crease. No more guessing."
)

divider()

# --- Build 8 ---
heading2("BUILD 8 — WhatsApp Integration")
body(
    "Same as Discord. Once Discord is stable and tested, the bot connects to WhatsApp groups "
    "through the same engine. Only the message connector changes. The AI underneath is identical."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — End Vision Flow
# ══════════════════════════════════════════════════════════════════════════════
heading1("SECTION 5 — The End Vision")
divider()

body("Once everything above is built, here is what the full system does — end to end:", bold=True)
doc.add_paragraph()

steps = [
    ("Step 1", "Someone types in Discord or WhatsApp:",
     '"Bought SPX 7570 call, premium 1.40"'),
    ("Step 2", "Bot reads the message in real time", ""),
    ("Step 3", "Understands:", "SPX | Call | Strike 7570 | Premium 1.40"),
    ("Step 4", "Checks today's full conversation for context", ""),
    ("Step 5", "Runs AI prediction using 3-agent committee vote", ""),
    ("Step 6", "Posts back in 10 seconds:",
     "AI Decision: HOLD | Confidence: 61%\nHistorical win rate for this setup: 43% | Proceed with caution"),
    ("Step 7", "Signal logged against Person A's account", ""),
    ("Step 8", "3 days later: trade outcome checked automatically", ""),
    ("Step 9", "Person A's accuracy score updated:", "14 signals, 71% correct"),
]

for step, title, detail in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{step}: ")
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        r3 = p2.add_run(detail)
        r3.italic = True; r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Full Roadmap Table
# ══════════════════════════════════════════════════════════════════════════════
heading1("SECTION 6 — Full Roadmap at a Glance")
divider()

add_table(
    ["Phase", "What We Build", "Why It Matters"],
    [
        ["Phase 1", "Fix AI accuracy — SPY + TSLA focus",          "Everything depends on the AI being trustworthy"],
        ["Phase 2", "Bulk backtesting — top 10 S&P 500",           "Prove the accuracy story with real data"],
        ["Phase 3", "Natural language parser",                       "Anyone can type anything — system understands"],
        ["Phase 4", "Discord bot — live group integration",         "First real-world deployment"],
        ["Phase 5", "Full day memory + reply chains",               "Bot understands the full conversation"],
        ["Phase 6", "Signal performance tracking / leaderboard",    "Know who to follow, who to ignore"],
        ["Phase 7", "WhatsApp integration",                         "Same bot, second channel"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Today vs After Full Build
# ══════════════════════════════════════════════════════════════════════════════
heading1("SECTION 7 — Where We Are Today vs Where We Are Going")
divider()

add_table(
    ["Today", "After Full Build"],
    [
        ["24% prediction accuracy",          "50%+ accuracy on SPY / TSLA"],
        ["Test one stock at a time manually", "Batch test 10 stocks overnight"],
        ["Must use strict form inputs",       "Type anything in plain English"],
        ["No Discord / WhatsApp connection",  "Bot reads groups in real time"],
        ["No conversation memory",            "Full day context always active"],
        ["No signal tracking",                "Live leaderboard per person and group"],
    ],
    header_color="1A56DB",
    alt_color="FEF9C3",
)

doc.add_paragraph()
divider()

# Final summary box
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("One Line Summary")
r.bold = True; r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

blockquote(
    "We are building a system that watches your trading groups, understands what everyone is saying, "
    "predicts whether each trade is good or bad, tracks who was right over time, "
    "and gets smarter with every single run."
)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("System status: 567 tests passing | Core engine stable | Phase 1 accuracy improvements in progress")
r2.italic = True; r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
